"""
Document generation utility for creating Word documents from workflow state.
"""

import os
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from pydantic import BaseModel, Field
from app.models.openai import llm_model
import logging

logger = logging.getLogger(__name__)

# Template path
TEMPLATE_PATH = Path(__file__).parent.parent / "doc_geneation" / "Report.docx"


class FIRPlaceholders(BaseModel):
    """Pydantic model for extracting FIR placeholders"""
    name_of_accused: str = Field(description="Full name of the accused person as mentioned in FIR")
    case_title: str = Field(description="Case title in format: STATE OF [STATE] vs. [ACCUSED NAME] with any qualifiers like (JUVENILE)")
    fir_date: str = Field(description="Date of FIR in format: DD.MM.YYYY or DD/MM/YYYY")
    sections_invoked: str = Field(description="All sections invoked in FIR, formatted as: Section 8(c), Section 20(b)(ii)(B), Section 29")


def extract_fir_placeholders(fir_facts: Dict[str, str], pdf_content: str = None) -> Dict[str, str]:
    """
    Extract placeholders from FIR facts using LLM with Pydantic model.
    
    Args:
        fir_facts: Dictionary of FIR facts
        pdf_content: Raw PDF content for better extraction
        
    Returns:
        Dictionary with placeholder values
    """
    try:
        # Combine FIR facts into a single text
        fir_text = ""
        if fir_facts:
            for key, value in fir_facts.items():
                fir_text += f"{key.replace('_', ' ').title()}: {value}\n"
        
        if pdf_content:
            fir_text = pdf_content[:2000]  # Use first 2000 chars for context
        
        if not fir_text:
            return {
                "name_of_accused": "Unknown",
                "case_title": "STATE OF GUJARAT vs. Unknown",
                "fir_date": "Not specified",
                "sections_invoked": "NDPS Act sections"
            }
        
        # Use LLM to extract placeholders
        prompt = f"""Extract the following information from the FIR content provided below:

1. **name_of_accused**: Full name of the accused person as mentioned in FIR
2. **case_title**: Format as "STATE OF GUJARAT vs. [ACCUSED NAME]" with qualifiers like "(JUVENILE)" if applicable
3. **fir_date**: Date of FIR in format DD.MM.YYYY or DD/MM/YYYY
4. **sections_invoked**: All sections invoked in FIR, formatted as: "Section 8(c), Section 20(b)(ii)(B), Section 29"

FIR Content:
{fir_text}

Extract the information accurately from the FIR content above.
"""
        
        result = llm_model.with_structured_output(FIRPlaceholders).invoke(prompt)
        
        return {
            "name_of_accused": result.name_of_accused,
            "case_title": result.case_title,
            "fir_date": result.fir_date,
            "sections_invoked": result.sections_invoked
        }
    except Exception as e:
        logger.error(f"Error extracting FIR placeholders: {e}", exc_info=True)
        # Fallback to basic extraction
        name_of_accused = "Unknown"
        if fir_facts and "identity_of_accused" in fir_facts:
            identity = fir_facts["identity_of_accused"]
            lines = identity.split(",")
            if lines:
                name_of_accused = lines[0].strip()
                for prefix in ["Name:", "Accused:", "The accused"]:
                    if name_of_accused.startswith(prefix):
                        name_of_accused = name_of_accused[len(prefix):].strip()
        
        return {
            "name_of_accused": name_of_accused,
            "case_title": f"STATE OF GUJARAT vs. {name_of_accused}",
            "fir_date": "Not specified",
            "sections_invoked": "NDPS Act sections"
        }


def add_formatted_section(doc: Document, title: str, content: str, level: int = 1):
    """
    Add a formatted section to the document with proper styling.
    
    Args:
        doc: Document object
        title: Section title
        content: Section content
        level: Heading level (1-3)
    """
    # Add heading
    heading = doc.add_heading(title, level=level)
    heading.style.font.size = Pt(14 - level) if level <= 2 else Pt(12)
    heading.style.font.bold = True
    
    # Add content paragraph
    if content:
        para = doc.add_paragraph(content)
        para.style.font.size = Pt(11)
        para.style.paragraph_format.space_after = Pt(6)


def format_section_content(doc: Document, formatted_state: Dict[str, Any]):
    """
    Format all available sections and add them to the document with proper formatting.
    
    Args:
        doc: Document object
        formatted_state: Already formatted state dictionary (from format_state_for_display)
    """
    """
    Format all available sections into a formatted string for the document.
    
    Args:
        formatted_state: Already formatted state dictionary (from format_state_for_display)
        
    Returns:
        Formatted string with all sections
    """
    # FIR Facts
    if formatted_state.get("fir_facts"):
        heading = doc.add_heading("FIR FACTS", level=1)
        heading.style.font.bold = True
        for key, value in formatted_state["fir_facts"].items():
            key_formatted = key.replace("_", " ").title()
            para = doc.add_paragraph()
            run1 = para.add_run(f"{key_formatted}: ")
            run1.bold = True
            run1.font.size = Pt(11)
            run2 = para.add_run(value)
            run2.font.size = Pt(11)
        doc.add_paragraph()  # Add spacing
    
    # NDPS Sections
    if formatted_state.get("ndps_sections"):
        heading = doc.add_heading("NDPS ACT SECTIONS", level=1)
        heading.style.font.bold = True
        for section in formatted_state["ndps_sections"]:
            if isinstance(section, dict):
                section_num = section.get("section_number", "")
                description = section.get("section_description", "")
                relevance = section.get("why_section_is_relevant", "")
                
                para = doc.add_paragraph()
                run1 = para.add_run(f"Section {section_num}: ")
                run1.bold = True
                run1.font.size = Pt(11)
                run2 = para.add_run(description)
                run2.font.size = Pt(11)
                
                if relevance:
                    para2 = doc.add_paragraph()
                    run3 = para2.add_run("  Relevance: ")
                    run3.italic = True
                    run3.font.size = Pt(10)
                    run4 = para2.add_run(relevance)
                    run4.font.size = Pt(10)
        doc.add_paragraph()
    
    # BNS Sections
    if formatted_state.get("bns_sections"):
        heading = doc.add_heading("BHARATIYA NYAYA SANHITA (BNS) SECTIONS", level=1)
        heading.style.font.bold = True
        for section in formatted_state["bns_sections"]:
            if isinstance(section, dict):
                section_num = section.get("section_number", "")
                description = section.get("section_description", "")
                relevance = section.get("why_section_is_relevant", "")
                
                para = doc.add_paragraph()
                run1 = para.add_run(f"Section {section_num}: ")
                run1.bold = True
                run1.font.size = Pt(11)
                run2 = para.add_run(description)
                run2.font.size = Pt(11)
                
                if relevance:
                    para2 = doc.add_paragraph()
                    run3 = para2.add_run("  Relevance: ")
                    run3.italic = True
                    run3.font.size = Pt(10)
                    run4 = para2.add_run(relevance)
                    run4.font.size = Pt(10)
        doc.add_paragraph()
    
    # BNSS Sections
    if formatted_state.get("bnss_sections"):
        heading = doc.add_heading("BHARATIYA NAGARIK SURAKSHA SANHITA (BNSS) SECTIONS", level=1)
        heading.style.font.bold = True
        for section in formatted_state["bnss_sections"]:
            if isinstance(section, dict):
                section_num = section.get("section_number", "")
                description = section.get("section_description", "")
                relevance = section.get("why_section_is_relevant", "")
                
                para = doc.add_paragraph()
                run1 = para.add_run(f"Section {section_num}: ")
                run1.bold = True
                run1.font.size = Pt(11)
                run2 = para.add_run(description)
                run2.font.size = Pt(11)
                
                if relevance:
                    para2 = doc.add_paragraph()
                    run3 = para2.add_run("  Relevance: ")
                    run3.italic = True
                    run3.font.size = Pt(10)
                    run4 = para2.add_run(relevance)
                    run4.font.size = Pt(10)
        doc.add_paragraph()
    
    # BSA Sections
    if formatted_state.get("bsa_sections"):
        heading = doc.add_heading("BHARATIYA SAKSHYA ADHINIYAM (BSA) SECTIONS", level=1)
        heading.style.font.bold = True
        for section in formatted_state["bsa_sections"]:
            if isinstance(section, dict):
                section_num = section.get("section_number", "")
                description = section.get("section_description", "")
                relevance = section.get("why_section_is_relevant", "")
                
                para = doc.add_paragraph()
                run1 = para.add_run(f"Section {section_num}: ")
                run1.bold = True
                run1.font.size = Pt(11)
                run2 = para.add_run(description)
                run2.font.size = Pt(11)
                
                if relevance:
                    para2 = doc.add_paragraph()
                    run3 = para2.add_run("  Relevance: ")
                    run3.italic = True
                    run3.font.size = Pt(10)
                    run4 = para2.add_run(relevance)
                    run4.font.size = Pt(10)
        doc.add_paragraph()
    
    # Investigation Plan
    if formatted_state.get("investigation_plan"):
        heading = doc.add_heading("INVESTIGATION PLAN", level=1)
        heading.style.font.bold = True
        investigation_plan = formatted_state["investigation_plan"]
        if isinstance(investigation_plan, list) and len(investigation_plan) > 0:
            for item in investigation_plan:
                if isinstance(item, dict):
                    title = item.get("title", "")
                    date_range = item.get("date_range", "")
                    description = item.get("description", "")
                    
                    if title or description:
                        para = doc.add_paragraph()
                        if title:
                            run1 = para.add_run(f"{title}")
                            run1.bold = True
                            run1.font.size = Pt(11)
                        if date_range:
                            run2 = para.add_run(f" ({date_range})")
                            run2.font.size = Pt(11)
                        if description:
                            run3 = para.add_run(f": {description}")
                            run3.font.size = Pt(11)
                elif isinstance(item, str):
                    # Handle case where item is a string
                    para = doc.add_paragraph(item)
                    para.style.font.size = Pt(11)
        doc.add_paragraph()
    
    # Timeline
    if formatted_state.get("investigation_and_legal_timeline"):
        timeline = formatted_state["investigation_and_legal_timeline"]
        heading = doc.add_heading("INVESTIGATION & LEGAL TIMELINE", level=1)
        heading.style.font.bold = True
        if isinstance(timeline, dict):
            date_str = timeline.get("date_string", "")
            timeline_str = timeline.get("timeline", "")
            
            para = doc.add_paragraph()
            run1 = para.add_run(f"Date: ")
            run1.bold = True
            run1.font.size = Pt(11)
            run2 = para.add_run(date_str)
            run2.font.size = Pt(11)
            
            para2 = doc.add_paragraph(timeline_str)
            para2.style.font.size = Pt(11)
        doc.add_paragraph()
    
    # Evidence Checklist
    if formatted_state.get("evidence_checklist"):
        heading = doc.add_heading("EVIDENCE CHECKLIST", level=1)
        heading.style.font.bold = True
        para = doc.add_paragraph(formatted_state['evidence_checklist'])
        para.style.font.size = Pt(11)
        doc.add_paragraph()
    
    # Do's and Don'ts
    if formatted_state.get("dos") or formatted_state.get("donts"):
        heading = doc.add_heading("DO'S AND DON'TS", level=1)
        heading.style.font.bold = True
        if formatted_state.get("dos"):
            para = doc.add_paragraph()
            run = para.add_run("DO'S:")
            run.bold = True
            run.font.size = Pt(11)
            for item in formatted_state["dos"]:
                para2 = doc.add_paragraph(f"  • {item}")
                para2.style.font.size = Pt(11)
        if formatted_state.get("donts"):
            para = doc.add_paragraph()
            run = para.add_run("DON'TS:")
            run.bold = True
            run.font.size = Pt(11)
            for item in formatted_state["donts"]:
                para2 = doc.add_paragraph(f"  • {item}")
                para2.style.font.size = Pt(11)
        doc.add_paragraph()
    
    # Prosecution Weaknesses
    if formatted_state.get("potential_prosecution_weaknesses"):
        heading = doc.add_heading("POTENTIAL PROSECUTION WEAKNESSES", level=1)
        heading.style.font.bold = True
        weaknesses = formatted_state["potential_prosecution_weaknesses"]
        if isinstance(weaknesses, dict):
            for key, value in weaknesses.items():
                para = doc.add_paragraph()
                run1 = para.add_run(f"{key}:")
                run1.bold = True
                run1.font.size = Pt(11)
                para2 = doc.add_paragraph(f"  {value}")
                para2.style.font.size = Pt(11)
        doc.add_paragraph()
    
    # Historical Cases
    if formatted_state.get("historical_cases"):
        heading = doc.add_heading("HISTORICAL CASES", level=1)
        heading.style.font.bold = True
        for case in formatted_state["historical_cases"]:
            if isinstance(case, dict):
                title = case.get("title", "")
                summary = case.get("summary", "")
                
                para = doc.add_paragraph()
                run = para.add_run(title)
                run.bold = True
                run.font.size = Pt(11)
                
                para2 = doc.add_paragraph(f"  {summary}")
                para2.style.font.size = Pt(11)
        doc.add_paragraph()
    
    # Defence Perspective & Rebuttal
    if formatted_state.get("defence_perspective_rebuttal"):
        heading = doc.add_heading("DEFENCE PERSPECTIVE & PROSECUTION REBUTTAL", level=1)
        heading.style.font.bold = True
        for item in formatted_state["defence_perspective_rebuttal"]:
            if isinstance(item, dict):
                defence = item.get("defence_perspective", [])
                rebuttal = item.get("rebuttal", [])
                if defence:
                    para = doc.add_paragraph()
                    run = para.add_run("Defence Perspective:")
                    run.bold = True
                    run.font.size = Pt(11)
                    for point in defence:
                        para2 = doc.add_paragraph(f"  • {point}")
                        para2.style.font.size = Pt(11)
                if rebuttal:
                    para = doc.add_paragraph()
                    run = para.add_run("Prosecution Rebuttal:")
                    run.bold = True
                    run.font.size = Pt(11)
                    for point in rebuttal:
                        para2 = doc.add_paragraph(f"  • {point}")
                        para2.style.font.size = Pt(11)
                doc.add_paragraph()
        doc.add_paragraph()
    
    # Summary for the Court
    if formatted_state.get("summary_for_the_court"):
        summary = formatted_state["summary_for_the_court"]
        heading = doc.add_heading("SUMMARY FOR THE COURT", level=1)
        heading.style.font.bold = True
        if isinstance(summary, dict):
            para = doc.add_paragraph()
            run1 = para.add_run("Case Title: ")
            run1.bold = True
            run1.font.size = Pt(11)
            run2 = para.add_run(summary.get('case_title', ''))
            run2.font.size = Pt(11)
            
            para2 = doc.add_paragraph()
            run3 = para2.add_run("NDPS Sections: ")
            run3.bold = True
            run3.font.size = Pt(11)
            run4 = para2.add_run(', '.join(summary.get('ndps_sections', [])))
            run4.font.size = Pt(11)
            
            heading2 = doc.add_heading("Core Issue", level=2)
            heading2.style.font.bold = True
            para3 = doc.add_paragraph(summary.get('core_issue', ''))
            para3.style.font.size = Pt(11)
            
            para4 = doc.add_paragraph()
            run5 = para4.add_run("Date & Place: ")
            run5.bold = True
            run5.font.size = Pt(11)
            run6 = para4.add_run(summary.get('date_and_place', ''))
            run6.font.size = Pt(11)
            
            para5 = doc.add_paragraph()
            run7 = para5.add_run("Recovery: ")
            run7.bold = True
            run7.font.size = Pt(11)
            run8 = para5.add_run(summary.get('recovery', ''))
            run8.font.size = Pt(11)
            
            para6 = doc.add_paragraph()
            run9 = para6.add_run("Quantity: ")
            run9.bold = True
            run9.font.size = Pt(11)
            run10 = para6.add_run(summary.get('quantity', ''))
            run10.font.size = Pt(11)
            
            safeguards = summary.get('safeguards', [])
            if safeguards:
                para7 = doc.add_paragraph()
                run11 = para7.add_run("Safeguards:")
                run11.bold = True
                run11.font.size = Pt(11)
                for safeguard in safeguards:
                    para8 = doc.add_paragraph(f"  ✓ {safeguard}", )
                    para8.style.font.size = Pt(11)
            
            heading2 = doc.add_heading("Why the Case is Strong", level=2)
            heading2.style.font.bold = True
            conscious = summary.get('conscious_possession_proven', [])
            if conscious:
                para9 = doc.add_paragraph()
                run12 = para9.add_run("Conscious Possession Proven:")
                run12.bold = True
                run12.font.size = Pt(11)
                for point in conscious:
                    para10 = doc.add_paragraph(f"  • {point}")
                    para10.style.font.size = Pt(11)
            procedural = summary.get('procedural_compliance', [])
            if procedural:
                para11 = doc.add_paragraph()
                run13 = para11.add_run("Procedural Compliance:")
                run13.bold = True
                run13.font.size = Pt(11)
                for point in procedural:
                    para12 = doc.add_paragraph(f"  • {point}")
                    para12.style.font.size = Pt(11)
            legal = summary.get('legal_position', [])
            if legal:
                para13 = doc.add_paragraph()
                run14 = para13.add_run("Legal Position:")
                run14.bold = True
                run14.font.size = Pt(11)
                for point in legal:
                    para14 = doc.add_paragraph(f"  • {point}")
                    para14.style.font.size = Pt(11)
            
            heading2 = doc.add_heading("Judicial Balance", level=2)
            heading2.style.font.bold = True
            para15 = doc.add_paragraph(summary.get('judicial_balance', ''))
            para15.style.font.size = Pt(11)
            
            prayers = summary.get('prosecution_prayer', [])
            if prayers:
                para16 = doc.add_paragraph()
                run15 = para16.add_run("Prosecution Prayer:")
                run15.bold = True
                run15.font.size = Pt(11)
                for prayer in prayers:
                    para17 = doc.add_paragraph(f"  • {prayer}", )
                    para17.style.font.size = Pt(11)
        doc.add_paragraph()
    
    # Chargesheet
    if formatted_state.get("chargesheet"):
        chargesheet = formatted_state["chargesheet"]
        heading = doc.add_heading("CHARGESHEET", level=1)
        heading.style.font.bold = True
        if isinstance(chargesheet, dict):
            para = doc.add_paragraph()
            run1 = para.add_run("Case Title: ")
            run1.bold = True
            run1.font.size = Pt(11)
            run2 = para.add_run(chargesheet.get('case_title', ''))
            run2.font.size = Pt(11)
            
            ndps = chargesheet.get('ndps_sections', [])
            bns = chargesheet.get('bns_sections', [])
            bnss = chargesheet.get('bnss_sections', [])
            bsa = chargesheet.get('bsa_sections', [])
            all_sections = []
            if ndps:
                all_sections.append(f"NDPS Act – {', '.join(ndps)}")
            if bns:
                all_sections.append(f"BNS – {', '.join(bns)}")
            if bnss:
                all_sections.append(f"BNSS – {', '.join(bnss)}")
            if bsa:
                all_sections.append(f"BSA – {', '.join(bsa)}")
            if all_sections:
                para2 = doc.add_paragraph(' | '.join(all_sections))
                para2.style.font.size = Pt(11)
            
            heading2 = doc.add_heading("Core Issue", level=2)
            heading2.style.font.bold = True
            para3 = doc.add_paragraph(chargesheet.get('core_issue', ''))
            para3.style.font.size = Pt(11)
            
            para4 = doc.add_paragraph()
            run3 = para4.add_run("Date & Place: ")
            run3.bold = True
            run3.font.size = Pt(11)
            run4 = para4.add_run(chargesheet.get('date_and_place', ''))
            run4.font.size = Pt(11)
            
            para5 = doc.add_paragraph()
            run5 = para5.add_run("Recovery: ")
            run5.bold = True
            run5.font.size = Pt(11)
            run6 = para5.add_run(chargesheet.get('recovery', ''))
            run6.font.size = Pt(11)
            
            para6 = doc.add_paragraph()
            run7 = para6.add_run("Quantity: ")
            run7.bold = True
            run7.font.size = Pt(11)
            run8 = para6.add_run(chargesheet.get('quantity', ''))
            run8.font.size = Pt(11)
            
            safeguards = chargesheet.get('safeguards', [])
            if safeguards:
                para7 = doc.add_paragraph()
                run9 = para7.add_run("Safeguards:")
                run9.bold = True
                run9.font.size = Pt(11)
                for safeguard in safeguards:
                    para8 = doc.add_paragraph(f"  ✓ {safeguard}")
                    para8.style.font.size = Pt(11)
            
            heading2 = doc.add_heading("Why the Case is Strong", level=2)
            heading2.style.font.bold = True
            conscious = chargesheet.get('conscious_possession_proven', [])
            if conscious:
                para9 = doc.add_paragraph()
                run10 = para9.add_run("Conscious Possession Proven:")
                run10.bold = True
                run10.font.size = Pt(11)
                for point in conscious:
                    para10 = doc.add_paragraph(f"  • {point}")
                    para10.style.font.size = Pt(11)
            procedural = chargesheet.get('procedural_compliance', [])
            if procedural:
                para11 = doc.add_paragraph()
                run11 = para11.add_run("Procedural Compliance:")
                run11.bold = True
                run11.font.size = Pt(11)
                for point in procedural:
                    para12 = doc.add_paragraph(f"  • {point}")
                    para12.style.font.size = Pt(11)
            legal = chargesheet.get('legal_position', [])
            if legal:
                para13 = doc.add_paragraph()
                run12 = para13.add_run("Legal Position:")
                run12.bold = True
                run12.font.size = Pt(11)
                for point in legal:
                    para14 = doc.add_paragraph(f"  • {point}")
                    para14.style.font.size = Pt(11)
            
            heading2 = doc.add_heading("Judicial Balance", level=2)
            heading2.style.font.bold = True
            para15 = doc.add_paragraph(chargesheet.get('judicial_balance', ''))
            para15.style.font.size = Pt(11)
            
            prayers = chargesheet.get('prosecution_prayer', [])
            if prayers:
                para16 = doc.add_paragraph()
                run13 = para16.add_run("Prosecution Prayer:")
                run13.bold = True
                run13.font.size = Pt(11)
                for prayer in prayers:
                    para17 = doc.add_paragraph(f"  • {prayer}")
                    para17.style.font.size = Pt(11)
        doc.add_paragraph()


def generate_document(workflow_state: Dict[str, Any]) -> bytes:
    """
    Generate a Word document from workflow state.
    
    Args:
        workflow_state: Complete workflow state dictionary
        
    Returns:
        Bytes of the generated document
    """
    try:
        # Load template
        if not TEMPLATE_PATH.exists():
            raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")
        
        doc = Document(str(TEMPLATE_PATH))
        
        # Extract FIR placeholders using LLM
        fir_facts = workflow_state.get("fir_facts", {})
        pdf_content = workflow_state.get("pdf_content_in_english", "")
        placeholders = extract_fir_placeholders(fir_facts, pdf_content)
        
        # Replace placeholders in all paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if "{{" in text:
                text = text.replace("{{name_of_accused}}", placeholders["name_of_accused"])
                text = text.replace("{{case_title}}", placeholders["case_title"])
                text = text.replace("{{fir_date}}", placeholders["fir_date"])
                text = text.replace("{{sections_invoked}}", placeholders["sections_invoked"])
                if text != paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run(text)
                    run.font.size = Pt(11)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if "{{" in text:
                            text = text.replace("{{name_of_accused}}", placeholders["name_of_accused"])
                            text = text.replace("{{case_title}}", placeholders["case_title"])
                            text = text.replace("{{fir_date}}", placeholders["fir_date"])
                            text = text.replace("{{sections_invoked}}", placeholders["sections_invoked"])
                            if text != paragraph.text:
                                paragraph.clear()
                                run = paragraph.add_run(text)
                                run.font.size = Pt(11)
        
        # Find and replace {{rest_content}} placeholder
        rest_content_found = False
        rest_content_paragraph = None
        for paragraph in doc.paragraphs:
            if "{{rest_content}}" in paragraph.text:
                rest_content_paragraph = paragraph
                rest_content_found = True
                break
        
        # Also check in tables
        if not rest_content_found:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if "{{rest_content}}" in paragraph.text:
                                rest_content_paragraph = paragraph
                                rest_content_found = True
                                break
                        if rest_content_found:
                            break
                    if rest_content_found:
                        break
                if rest_content_found:
                    break
        
        # Format and add all sections
        from ..routes.utils import format_state_for_display
        formatted_state = format_state_for_display(workflow_state)
        
        if rest_content_paragraph:
            # Clear the placeholder paragraph
            rest_content_paragraph.clear()
            # Add formatted content starting from this paragraph
            format_section_content(doc, formatted_state)
        else:
            # If placeholder not found, add content at the end
            doc.add_page_break()
            format_section_content(doc, formatted_state)
        
        # Save to bytes
        from io import BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()
        
    except Exception as e:
        logger.error(f"Error generating document: {e}", exc_info=True)
        raise


